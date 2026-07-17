from docx import Document

def create_resume(name, email, phone, education, skills, projects):

    doc = Document()

    doc.add_heading(name, 0)

    doc.add_paragraph(f"Email: {email}")
    doc.add_paragraph(f"Phone: {phone}")

    doc.add_heading("Education", level=1)
    doc.add_paragraph(education)

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(skills)

    doc.add_heading("Projects", level=1)
    doc.add_paragraph(projects)

    doc.save(f"resumes/{name}.docx")
